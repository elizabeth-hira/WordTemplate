import eel
import pandas as pd
from docx import Document
import os
import docx
from docx.shared import RGBColor

import main

eel.init("WEB")

merge_fields = ["«nickname»", "«hoursOfPlaying»", "«level»"]
nickname = "elizabeth_hira_"
hoursOfPlaying = 310
level = 58
fields = [nickname, hoursOfPlaying, level]
colors = [[0, 0, 0], [68, 114, 196], [0, 0, 0], [192, 0, 0]]


df = pd.read_excel("table.xlsx")
df['Begin date'] = pd.to_datetime(df['Begin date'])
short_df = pd.DataFrame(columns=df.columns)
print(df)




def add_run_copy(paragraph, run, text, rgb):
    r = paragraph.add_run(text=run.text if text is None else text, style=run.style)
    r.bold = run.bold
    r.italic = run.italic
    r.underline = run.underline
    r.font.all_caps = run.font.all_caps
    r.font.bold = run.font.bold
    r.font.color.rgb = run.font.color.rgb
    r.font.color.theme_color = run.font.color.theme_color
    #r.font.color.type = run.font.color.type
    r.font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])
    r.font.complex_script = run.font.complex_script
    r.font.cs_bold = run.font.cs_bold
    r.font.cs_italic = run.font.cs_italic
    r.font.double_strike = run.font.double_strike
    r.font.emboss = run.font.emboss
    r.font.hidden = run.font.hidden
    r.font.highlight_color = run.font.highlight_color
    r.font.imprint = run.font.imprint
    r.font.italic = run.font.italic
    r.font.math = run.font.math
    r.font.name = run.font.name
    r.font.no_proof = run.font.no_proof
    r.font.outline = run.font.outline
    r.font.rtl = run.font.rtl
    r.font.shadow = run.font.shadow
    r.font.size = run.font.size
    r.font.small_caps = run.font.small_caps
    r.font.snap_to_grid = run.font.snap_to_grid
    r.font.spec_vanish = run.font.spec_vanish
    r.font.strike = run.font.strike
    r.font.subscript = run.font.subscript
    r.font.superscript = run.font.superscript
    r.font.underline = run.font.underline
    r.font.web_hidden = run.font.web_hidden
    return r


@eel.expose
def get_short_table(num):
    try:
        print(num)
        num = int(num)
    except:
        return "wrong params"
    main.short_df = df.sample(n=num)
    return get_table(main.short_df)



@eel.expose
def get_table(dataframe):

    table_str = """<table border="1">
       <caption>Таблица</caption>
   <tr>
    <th>Mission</th>
    <th>Next Step</th>
    <th>Begin Date</th>
    <th>Percent of completion</th>
   </tr>
    """

    for index, row in dataframe.iterrows():
        table_str = table_str + "<tr><td>" + row['Mission'] + "</td><td>" + row['Next step'] + "</td><td>"\
                    + str(row['Begin date'].day) + "-" + str(row['Begin date'].month) + "-" +\
                    str(row['Begin date'].year) + "</td><td>" + str(row['Percent of completion']) + "</td></tr>"

    table_str = table_str + "</table>"
    return table_str

@eel.expose
def sort_table(column, ascending):
    print(column)
    print(ascending)
    print(short_df.sort_values(by=[column], ascending=ascending))
    short_df.sort_values(by=[column], ascending=ascending, inplace=True)
    return get_table(short_df)

@eel.expose
def add_row(mission, step, date, perc):
    print(mission)
    print(step)
    print(date)
    print(perc)
    row = pd.DataFrame([[mission, step, pd.to_datetime(date), int(perc)]], columns=df.columns)
    main.short_df = main.short_df.append(row, ignore_index=True)
    return get_table(main.short_df)


@eel.expose
def open_file(filename):
    print(filename)
    arr = filename.split("\\")
    name = arr[len(arr) - 1]
    print(name)
    doc = docx.Document(name)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for i in range(len(merge_fields)):
                if merge_fields[i] in run.text:
                    run.text = run.text.replace(merge_fields[i], str(fields[i]))
    for table in doc.tables:
        alignment = []
        runs = []
        num_row = 0
        for row in table.rows:
            if row.cells[0].text == "«firstCell»":
                for j in range(len(short_df.columns)):
                    row.cells[j].paragraphs[0].runs[0].text = str(short_df.iloc[0, j])

                    runs.append(row.cells[j].paragraphs[0].runs[0])
                    alignment.append(row.cells[j].paragraphs[0].alignment)
                    for run in range(1, len(row.cells[j].paragraphs[0].runs)):
                        row.cells[j].paragraphs[0].runs[run].text = ""
                num_row +=1
                for i in range(num_row, len(short_df)):
                    table.add_row()
                    row = table.rows[i]
                    for j in range(len(short_df.columns)):
                        add_run_copy(row.cells[j].paragraphs[0], runs[j], str(short_df.iloc[i, j]), colors[j])
                        row.cells[j].paragraphs[0].alignment = alignment[j]
                    num_row +=1
                break
            num_row +=1

    doc.add_picture('witcher.jpeg', width = docx.shared.Cm(16))
    doc.save('new_file.docx')



    os.startfile('new_file.docx')


eel.start("main.html")


